import requests
from bs4 import BeautifulSoup
import docx
import cloudscraper

#bypass Cloudfare
scraper = cloudscraper.create_scraper()  # returns a CloudScraper instance
base_url = "https://www.mtlnovel.com/i-just-want-to-be-a-shenhao-quietly/chapter-"

### Create a new docx document
##document = docx.Document()

# Iterate over the chapters
for i in range(11, 21):
    #creates document
    document = docx.Document()
    # Construct the URL of the chapter
    base_url = "https://www.mtlnovel.com/i-just-want-to-be-a-shenhao-quietly/chapter-"
    url = base_url + str(i) + "/"

    # Make a request to the website and get the HTML
    html = scraper.get(url).text

    # Parse the HTML with Beautiful Soup
    soup = BeautifulSoup(html, "html5lib")

    # Get the chapter title and text
    title = soup.h1
    div = soup.find_all("div", {"class":"par fontsize-16"})

    # Add the title and text to the document and prints them to monitor progress
    document.add_heading(title, level=1)
    print(title.getText() + " has been extracted.")
    for res in div:
        document.add_paragraph(res.text)

    # Save the chapters per document
    doc_name = "i-just-want-to-be-a-shenhao-quietly-chapter-"
    doc = doc_name + str(i) + ".docx"
    document.save(doc)

### Save the document
##document.save("novel.docx")
